"""
Módulo para generar reportes en diferentes formatos
"""
import json
from datetime import datetime
from typing import Dict, Any, List
from pathlib import Path
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Para generar gráficos sin interfaz gráfica
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from catalogos_rips import FINALIDADES_TECNOLOGIA_SALUD, CAUSAS_ATENCION


class GeneradorReportes:
    """Clase para generar reportes de análisis RIPS"""

    def __init__(self, directorio_salida: str = "output", nombre_archivo_json: str = None):
        """
        Inicializa el generador de reportes

        Args:
            directorio_salida: Directorio donde se guardarán los reportes
            nombre_archivo_json: Nombre del archivo JSON procesado (para crear subcarpeta)
        """
        self.directorio_salida = Path(directorio_salida)
        self.directorio_salida.mkdir(exist_ok=True)

        # Crear subcarpeta específica para este reporte si se proporciona nombre de archivo
        if nombre_archivo_json:
            # Extraer nombre base del archivo sin extensión
            nombre_base = Path(nombre_archivo_json).stem
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            self.directorio_reporte = self.directorio_salida / f"{nombre_base}_{timestamp}"
            self.directorio_reporte.mkdir(exist_ok=True)
        else:
            self.directorio_reporte = self.directorio_salida

    def generar_json(self, datos: Dict[str, Any], nombre_archivo: str) -> str:
        """
        Genera un reporte en formato JSON

        Args:
            datos: Datos a exportar
            nombre_archivo: Nombre del archivo (sin extensión)

        Returns:
            Ruta del archivo generado
        """
        ruta = self.directorio_reporte / f"{nombre_archivo}.json"

        with open(ruta, 'w', encoding='utf-8') as f:
            json.dump(datos, f, indent=2, ensure_ascii=False)

        return str(ruta)

    def generar_reporte_texto(self, analisis: Dict[str, Any], nombre_archivo: str) -> str:
        """
        Genera un reporte legible en formato texto

        Args:
            analisis: Diccionario con análisis completo
            nombre_archivo: Nombre del archivo (sin extensión)

        Returns:
            Ruta del archivo generado
        """
        ruta = self.directorio_reporte / f"{nombre_archivo}.txt"

        with open(ruta, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("REPORTE DE ANÁLISIS RIPS - RESOLUCIÓN 2275 DE 2023\n")
            f.write("=" * 80 + "\n\n")
            f.write(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

            # Información general
            info_gen = analisis.get("informacion_general", {})
            f.write("INFORMACIÓN GENERAL\n")
            f.write("-" * 80 + "\n")
            f.write(f"Número de Factura: {info_gen.get('num_factura', 'N/A')}\n")
            f.write(f"Documento Obligado: {info_gen.get('num_documento_obligado', 'N/A')}\n")
            f.write(f"Tipo de Nota: {info_gen.get('tipo_nota', 'N/A')}\n")
            f.write(f"Número de Nota: {info_gen.get('num_nota', 'N/A')}\n\n")

            # Análisis demográfico
            demo = analisis.get("analisis_demografico", {})
            f.write("ANÁLISIS DEMOGRÁFICO\n")
            f.write("-" * 80 + "\n")
            f.write(f"Total de Usuarios: {demo.get('total_usuarios', 0)}\n\n")

            f.write("Distribución por Sexo:\n")
            for sexo, cantidad in demo.get("distribucion_sexo", {}).items():
                f.write(f"  {sexo}: {cantidad}\n")

            f.write("\nDistribución por Ruta de Vida (Resolución 3280 de 2018):\n")
            for grupo, cantidad in demo.get("distribucion_ruta_vida", {}).items():
                f.write(f"  {grupo}: {cantidad}\n")

            f.write("\nDistribución por Tipo de Documento:\n")
            for tipo, cantidad in demo.get("distribucion_tipo_documento", {}).items():
                f.write(f"  {tipo}: {cantidad}\n")

            f.write("\nDistribución por Zona Territorial:\n")
            for zona, cantidad in demo.get("distribucion_zona_territorial", {}).items():
                zona_desc = "Urbana" if zona == "01" else "Rural" if zona == "02" else zona
                f.write(f"  {zona_desc}: {cantidad}\n")

            # Análisis de diagnósticos
            diag = analisis.get("analisis_diagnosticos", {})
            f.write("\n" + "=" * 80 + "\n")
            f.write("ANÁLISIS DE DIAGNÓSTICOS\n")
            f.write("=" * 80 + "\n")
            f.write(f"Total de Diagnósticos Registrados: {diag.get('total_diagnosticos', 0)}\n")
            f.write(f"Diagnósticos Únicos: {diag.get('diagnosticos_unicos', 0)}\n")
            f.write(f"Casos con Comorbilidades: {diag.get('diagnosticos_con_comorbilidades', 0)}\n\n")

            f.write("TODOS LOS DIAGNÓSTICOS IDENTIFICADOS:\n")
            f.write("-" * 80 + "\n")
            for i, diag_info in enumerate(diag.get("todos_diagnosticos", []), 1):
                f.write(f"{i}. [{diag_info.get('codigo')}] {diag_info.get('nombre')}\n")
                f.write(f"   Descripción: {diag_info.get('descripcion')}\n")
                f.write(f"   Cantidad: {diag_info.get('cantidad')} ({diag_info.get('porcentaje', 0):.1f}%)\n\n")

            # Comorbilidades
            comorbilidades = diag.get("comorbilidades_detalle", [])
            if comorbilidades:
                f.write("\nCOMORBILIDADES IDENTIFICADAS:\n")
                f.write("-" * 80 + "\n")
                for i, combo in enumerate(comorbilidades[:10], 1):  # Mostrar solo 10
                    info_principal = combo.get("info_principal", {})
                    f.write(f"{i}. Diagnóstico Principal: [{info_principal.get('codigo')}] {info_principal.get('nombre')}\n")
                    f.write("   Diagnósticos Relacionados:\n")
                    for rel in combo.get("relacionados", []):
                        f.write(f"   - [{rel.get('codigo')}] {rel.get('nombre')}\n")
                    f.write("\n")

            # Análisis de servicios
            serv = analisis.get("analisis_servicios", {})
            f.write("\n" + "=" * 80 + "\n")
            f.write("ANÁLISIS DE SERVICIOS\n")
            f.write("=" * 80 + "\n")
            f.write(f"Total de Consultas: {serv.get('total_consultas', 0)}\n")
            f.write(f"Servicios sin Autorización: {serv.get('servicios_sin_autorizacion', 0)}\n")
            f.write(f"Valor Total de Servicios: ${serv.get('total_valor_servicios', 0):,.2f}\n")
            f.write(f"Total Copagos: ${serv.get('total_copagos', 0):,.2f}\n\n")

            f.write("Distribución por Finalidad de Atención:\n")
            for finalidad, cantidad in serv.get("distribucion_finalidad", {}).items():
                finalidad_desc = self._obtener_descripcion_finalidad(finalidad)
                f.write(f"  {finalidad} - {finalidad_desc}: {cantidad}\n")

            f.write("\nDistribución por Causa de Atención:\n")
            for causa, cantidad in serv.get("distribucion_causa_atencion", {}).items():
                causa_desc = self._obtener_descripcion_causa(causa)
                f.write(f"  {causa} - {causa_desc}: {cantidad}\n")

            # Usuarios con múltiples consultas
            multiples = serv.get("usuarios_multiples_consultas", [])
            if multiples:
                f.write("\n" + "-" * 80 + "\n")
                f.write(f"USUARIOS CON MÚLTIPLES CONSULTAS: {len(multiples)}\n")
                f.write("-" * 80 + "\n")
                for user in multiples[:10]:  # Mostrar solo 10
                    f.write(f"Documento: {user.get('tipo_documento')} {user.get('documento')}\n")
                    f.write(f"Cantidad de Consultas: {user.get('cantidad_consultas')}\n")
                    f.write(f"Fechas: {', '.join(user.get('fechas', []))}\n\n")

            # Análisis de acudientes
            acud = analisis.get("analisis_acudientes", {})
            f.write("\n" + "=" * 80 + "\n")
            f.write("ANÁLISIS DE ACUDIENTES\n")
            f.write("=" * 80 + "\n")
            f.write(f"Total de Acudientes Identificados: {acud.get('total_acudientes', 0)}\n")
            f.write(f"Menores con Acudiente: {acud.get('total_menores_con_acudiente', 0)}\n\n")

            f.write("=" * 80 + "\n")
            f.write("FIN DEL REPORTE\n")
            f.write("=" * 80 + "\n")

        return str(ruta)

    def generar_reporte_csv(self, analisis: Dict[str, Any], nombre_archivo: str) -> str:
        """
        Genera un reporte CSV con diagnósticos detallados

        Args:
            analisis: Diccionario con análisis completo
            nombre_archivo: Nombre del archivo (sin extensión)

        Returns:
            Ruta del archivo generado
        """
        ruta = self.directorio_reporte / f"{nombre_archivo}.csv"

        diag = analisis.get("analisis_diagnosticos", {})
        todos_diagnosticos = diag.get("todos_diagnosticos", [])

        with open(ruta, 'w', encoding='utf-8') as f:
            f.write("Codigo,Nombre,Descripcion,Cantidad,Porcentaje\n")
            for diag_info in todos_diagnosticos:
                codigo = diag_info.get('codigo', '')
                nombre = diag_info.get('nombre', '').replace(',', ';')
                descripcion = diag_info.get('descripcion', '').replace(',', ';')
                cantidad = diag_info.get('cantidad', 0)
                porcentaje = diag_info.get('porcentaje', 0)
                f.write(f'"{codigo}","{nombre}","{descripcion}",{cantidad},{porcentaje:.1f}\n')

        return str(ruta)

    def _obtener_descripcion_finalidad(self, codigo: str) -> str:
        """Obtiene descripción de código de finalidad"""
        return FINALIDADES_TECNOLOGIA_SALUD.get(codigo, "No especificado")

    def _obtener_descripcion_causa(self, codigo: str) -> str:
        """Obtiene descripción de código de causa"""
        return CAUSAS_ATENCION.get(codigo, "No especificado")

    def generar_excel_completo(self, analisis: Dict[str, Any], datos_rips: Dict[str, Any], validacion: Dict[str, Any], nombre_archivo: str, analisis_avanzado: Dict[str, Any] = None) -> str:
        """
        Genera un archivo Excel completo con múltiples hojas de análisis

        Args:
            analisis: Diccionario con análisis completo
            datos_rips: Datos originales RIPS
            validacion: Diccionario con validación de calidad
            nombre_archivo: Nombre del archivo (sin extensión)
            analisis_avanzado: Diccionario con análisis avanzado (opcional)

        Returns:
            Ruta del archivo generado
        """
        ruta = self.directorio_reporte / f"{nombre_archivo}.xlsx"

        wb = Workbook()
        # Eliminar hoja por defecto
        wb.remove(wb.active)

        # 1. Hoja de Información General
        self._crear_hoja_info_general(wb, analisis)

        # 2. Hoja de Datos Personales de Pacientes
        self._crear_hoja_datos_personales(wb, datos_rips)

        # 3. Hoja de Servicios por Paciente
        self._crear_hoja_servicios_paciente(wb, datos_rips)

        # 4. Hoja de Análisis Demográfico
        self._crear_hoja_demografico(wb, analisis)

        # 5. Hoja de Diagnósticos
        self._crear_hoja_diagnosticos(wb, analisis)

        # 6. Hoja de Población Gestante
        self._crear_hoja_poblacion_gestante(wb, analisis)

        # 7. Hoja de Distribución por Ruta de Vida
        self._crear_hoja_distribucion_edad(wb, analisis)

        # 8. Hoja de Distribución Territorial
        self._crear_hoja_distribucion_territorial(wb, analisis)

        # 9. Hoja de Análisis de Servicios
        self._crear_hoja_servicios(wb, analisis)

        # 10. Hoja de Acudientes
        self._crear_hoja_acudientes(wb, analisis)

        # 11. Hoja de Validación de Calidad
        self._crear_hoja_validacion_calidad(wb, validacion)

        # 12. Hoja de Alertas Automáticas (si existe análisis avanzado)
        if analisis_avanzado:
            self._crear_hoja_alertas(wb, analisis_avanzado)

        # 13. Hoja de Indicadores de Calidad (si existe análisis avanzado)
        if analisis_avanzado:
            self._crear_hoja_indicadores_calidad(wb, analisis_avanzado)

        # 14. Hoja de Análisis de Morbilidad (si existe análisis avanzado)
        if analisis_avanzado:
            self._crear_hoja_morbilidad_especifica(wb, analisis_avanzado)

        # 15. Hoja de Procedimientos CUPS
        self._crear_hoja_procedimientos(wb, analisis)

        # 16. Hoja de Tipo de Usuario/Régimen
        self._crear_hoja_tipo_usuario(wb, analisis)

        # 17. Hoja de Incapacidades
        self._crear_hoja_incapacidades(wb, analisis)

        # 18. Hoja de Modalidad y Tipo Diagnóstico
        self._crear_hoja_modalidad_diagnostico(wb, analisis)

        # 19. Hoja de Prestadores
        self._crear_hoja_prestadores(wb, analisis)

        wb.save(ruta)
        return str(ruta)

    def _aplicar_estilo_encabezado(self, cell):
        """Aplica estilo de encabezado a una celda"""
        cell.font = Font(bold=True, color="FFFFFF", size=11)
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def _ajustar_columnas(self, ws, min_width=10, max_width=50):
        """Ajusta el ancho de las columnas automáticamente"""
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max(max_length + 2, min_width), max_width)
            ws.column_dimensions[column_letter].width = adjusted_width

    def _crear_hoja_info_general(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de información general"""
        ws = wb.create_sheet("Información General")
        info_gen = analisis.get("informacion_general", {})

        # Título
        ws['A1'] = "INFORMACIÓN GENERAL DEL REPORTE"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:B1')

        # Datos
        datos = [
            ["Número de Factura", info_gen.get('num_factura', 'N/A')],
            ["Documento Obligado", info_gen.get('num_documento_obligado', 'N/A')],
            ["Tipo de Nota", info_gen.get('tipo_nota', 'N/A')],
            ["Número de Nota", info_gen.get('num_nota', 'N/A')],
            ["Fecha de Generación", datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
        ]

        for i, (campo, valor) in enumerate(datos, start=3):
            ws[f'A{i}'] = campo
            ws[f'A{i}'].font = Font(bold=True)
            ws[f'B{i}'] = valor

        self._ajustar_columnas(ws)

    def _crear_hoja_datos_personales(self, wb: Workbook, datos_rips: Dict[str, Any]):
        """Crea hoja con datos personales de pacientes"""
        ws = wb.create_sheet("Datos Personales")

        # Encabezados
        encabezados = ["Tipo Doc", "Número Documento", "Sexo", "Fecha Nacimiento", "Edad",
                      "Municipio", "Zona Territorial", "Tipo Usuario", "País Residencia"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=1, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        # Extraer usuarios desde la estructura correcta
        from cargador_rips import CargadorRIPS
        cargador = CargadorRIPS()
        usuarios = datos_rips.get("usuarios", [])

        fila = 2
        for usuario in usuarios:
            ws.cell(row=fila, column=1, value=usuario.get("tipoDocumentoIdentificacion"))
            ws.cell(row=fila, column=2, value=usuario.get("numDocumentoIdentificacion"))
            ws.cell(row=fila, column=3, value="Femenino" if usuario.get("codSexo") == "F" else "Masculino" if usuario.get("codSexo") == "M" else usuario.get("codSexo"))
            ws.cell(row=fila, column=4, value=usuario.get("fechaNacimiento"))

            # Calcular edad
            fecha_nac = usuario.get("fechaNacimiento")
            edad = cargador.calcular_edad(fecha_nac) if fecha_nac else None
            ws.cell(row=fila, column=5, value=edad)

            ws.cell(row=fila, column=6, value=usuario.get("codMunicipioResidencia"))

            zona = usuario.get("codZonaTerritorialResidencia")
            zona_desc = "Urbana" if zona == "01" else "Rural" if zona == "02" else zona
            ws.cell(row=fila, column=7, value=zona_desc)

            ws.cell(row=fila, column=8, value=usuario.get("tipoUsuario"))
            ws.cell(row=fila, column=9, value=usuario.get("codPaisResidencia"))
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_servicios_paciente(self, wb: Workbook, datos_rips: Dict[str, Any]):
        """Crea hoja con servicios de ingreso por paciente"""
        ws = wb.create_sheet("Servicios por Paciente")

        # Encabezados
        encabezados = ["Tipo Doc Usuario", "Núm. Doc Usuario", "Sexo", "Fecha Consulta", "Modalidad",
                      "Finalidad", "Motivo Consulta", "Diagnóstico Principal", "Valor Servicio", "Copago"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=1, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        # Extraer consultas desde la estructura correcta
        usuarios = datos_rips.get("usuarios", [])
        fila = 2

        for usuario in usuarios:
            tipo_doc_usuario = usuario.get("tipoDocumentoIdentificacion")
            num_doc_usuario = usuario.get("numDocumentoIdentificacion")
            sexo = "Femenino" if usuario.get("codSexo") == "F" else "Masculino" if usuario.get("codSexo") == "M" else usuario.get("codSexo")

            consultas = usuario.get("servicios", {}).get("consultas", [])

            for consulta in consultas:
                ws.cell(row=fila, column=1, value=tipo_doc_usuario)
                ws.cell(row=fila, column=2, value=num_doc_usuario)
                ws.cell(row=fila, column=3, value=sexo)
                ws.cell(row=fila, column=4, value=consulta.get("fechaInicioAtencion"))
                ws.cell(row=fila, column=5, value=consulta.get("modalidadGrupoServicioTecSal"))
                ws.cell(row=fila, column=6, value=self._obtener_descripcion_finalidad(consulta.get("finalidadTecnologiaSalud", "")))
                ws.cell(row=fila, column=7, value=self._obtener_descripcion_causa(consulta.get("causaMotivoAtencion", "")))
                ws.cell(row=fila, column=8, value=consulta.get("codDiagnosticoPrincipal"))
                ws.cell(row=fila, column=9, value=consulta.get("vrServicio", 0))
                ws.cell(row=fila, column=10, value=consulta.get("valorPagoModerador", 0))
                fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_demografico(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis demográfico"""
        ws = wb.create_sheet("Análisis Demográfico")
        demo = analisis.get("analisis_demografico", {})

        # Resumen
        ws['A1'] = "ANÁLISIS DEMOGRÁFICO"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:C1')

        ws['A3'] = "Total de Usuarios:"
        ws['A3'].font = Font(bold=True)
        ws['B3'] = demo.get('total_usuarios', 0)

        # Distribución por Sexo
        ws['A5'] = "Distribución por Sexo"
        ws['A5'].font = Font(bold=True, size=12)
        ws['A6'] = "Sexo"
        ws['B6'] = "Cantidad"
        self._aplicar_estilo_encabezado(ws['A6'])
        self._aplicar_estilo_encabezado(ws['B6'])

        fila = 7
        for sexo, cantidad in demo.get("distribucion_sexo", {}).items():
            ws[f'A{fila}'] = "Femenino" if sexo == "F" else "Masculino"
            ws[f'B{fila}'] = cantidad
            fila += 1

        # Distribución por Ruta de Vida (Resolución 3280 de 2018)
        fila += 1
        ws[f'A{fila}'] = "Distribución por Ruta de Vida (Resolución 3280 de 2018)"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1
        ws[f'A{fila}'] = "Ruta de Vida"
        ws[f'B{fila}'] = "Cantidad"
        self._aplicar_estilo_encabezado(ws[f'A{fila}'])
        self._aplicar_estilo_encabezado(ws[f'B{fila}'])
        fila += 1

        for ruta_vida, cantidad in demo.get("distribucion_ruta_vida", {}).items():
            ws[f'A{fila}'] = ruta_vida
            ws[f'B{fila}'] = cantidad
            fila += 1

        # Distribución por Tipo de Documento
        fila += 1
        ws[f'A{fila}'] = "Distribución por Tipo de Documento"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1
        ws[f'A{fila}'] = "Tipo Documento"
        ws[f'B{fila}'] = "Cantidad"
        self._aplicar_estilo_encabezado(ws[f'A{fila}'])
        self._aplicar_estilo_encabezado(ws[f'B{fila}'])
        fila += 1

        for tipo, cantidad in demo.get("distribucion_tipo_documento", {}).items():
            ws[f'A{fila}'] = tipo
            ws[f'B{fila}'] = cantidad
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_diagnosticos(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de diagnósticos"""
        ws = wb.create_sheet("Diagnósticos")
        diag = analisis.get("analisis_diagnosticos", {})

        # Encabezados
        ws['A1'] = "TODOS LOS DIAGNÓSTICOS IDENTIFICADOS"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:E1')

        ws['A3'] = f"Total de Diagnósticos: {diag.get('total_diagnosticos', 0)}"
        ws['A4'] = f"Diagnósticos Únicos: {diag.get('diagnosticos_unicos', 0)}"

        encabezados = ["Código", "Nombre", "Descripción", "Cantidad", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=6, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        fila = 7
        for diag_info in diag.get("todos_diagnosticos", []):
            ws.cell(row=fila, column=1, value=diag_info.get('codigo'))
            ws.cell(row=fila, column=2, value=diag_info.get('nombre'))
            ws.cell(row=fila, column=3, value=diag_info.get('descripcion'))
            ws.cell(row=fila, column=4, value=diag_info.get('cantidad'))
            ws.cell(row=fila, column=5, value=f"{diag_info.get('porcentaje', 0):.1f}%")
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_poblacion_gestante(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis de población gestante"""
        ws = wb.create_sheet("Población Gestante")

        gestantes_data = analisis.get("analisis_poblacion_gestante", {})

        # Encabezado
        ws['A1'] = "ANÁLISIS DE POBLACIÓN GESTANTE"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="E91E63", end_color="E91E63", fill_type="solid")
        ws.merge_cells('A1:F1')

        # Resumen
        ws['A3'] = "RESUMEN"
        ws['A3'].font = Font(bold=True, size=12)
        ws.merge_cells('A3:C3')

        datos_resumen = [
            ["Total de Gestantes", gestantes_data.get('total_gestantes', 0)],
            ["Edad Promedio", f"{gestantes_data.get('edad_promedio', 0)} años"],
            ["Gestantes Adolescentes", gestantes_data.get('gestantes_adolescentes', 0)],
            ["Gestantes de Alto Riesgo", f"{gestantes_data.get('gestantes_alto_riesgo', 0)} ({gestantes_data.get('porcentaje_alto_riesgo', 0):.1f}%)"],
            ["Gestantes con Complicaciones", f"{gestantes_data.get('gestantes_con_complicaciones', 0)} ({gestantes_data.get('porcentaje_complicaciones', 0):.1f}%)"],
            ["Cobertura de Controles Prenatales", f"{gestantes_data.get('cobertura_controles_prenatales', 0):.1f}%"]
        ]

        fila = 4
        for campo, valor in datos_resumen:
            ws[f'A{fila}'] = campo
            ws[f'A{fila}'].font = Font(bold=True)
            ws[f'B{fila}'] = valor
            fila += 1

        # Distribución de Controles Prenatales
        fila += 2
        ws[f'A{fila}'] = "Distribución de Controles Prenatales"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{fila}:C{fila}')
        fila += 1

        encabezados = ["Número de Controles", "Cantidad de Gestantes", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        dist_controles = gestantes_data.get("distribucion_controles_prenatales", {})
        total_gestantes = gestantes_data.get('total_gestantes', 0)
        for num_controles, cantidad in sorted(dist_controles.items()):
            porcentaje = (cantidad / total_gestantes * 100) if total_gestantes > 0 else 0
            ws.cell(row=fila, column=1, value=f"{num_controles} controles")
            ws.cell(row=fila, column=2, value=cantidad)
            ws.cell(row=fila, column=3, value=f"{porcentaje:.1f}%")
            fila += 1

        # Distribución por edad
        fila += 2
        ws[f'A{fila}'] = "Distribución por Grupo de Edad"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        encabezados = ["Grupo de Edad", "Cantidad", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        dist_edad = gestantes_data.get("distribucion_edad_gestantes", {})
        for grupo, cantidad in dist_edad.items():
            porcentaje = (cantidad / total_gestantes * 100) if total_gestantes > 0 else 0
            ws.cell(row=fila, column=1, value=grupo)
            ws.cell(row=fila, column=2, value=cantidad)
            ws.cell(row=fila, column=3, value=f"{porcentaje:.1f}%")
            fila += 1

        # Top diagnósticos en gestantes
        fila += 2
        ws[f'A{fila}'] = "Top 10 Diagnósticos en Población Gestante"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        encabezados = ["Código", "Nombre", "Descripción", "Cantidad"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for diag_info in gestantes_data.get("top_diagnosticos_gestantes", []):
            ws.cell(row=fila, column=1, value=diag_info.get('codigo'))
            ws.cell(row=fila, column=2, value=diag_info.get('nombre'))
            ws.cell(row=fila, column=3, value=diag_info.get('descripcion'))
            ws.cell(row=fila, column=4, value=diag_info.get('cantidad'))
            fila += 1

        # Detalle de gestantes
        fila += 2
        ws[f'A{fila}'] = "Detalle de Gestantes Identificadas"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        encabezados = ["Tipo Doc", "Número Documento", "Edad", "Municipio", "Zona"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for gestante in gestantes_data.get("gestantes_detalle", []):
            ws.cell(row=fila, column=1, value=gestante.get("tipo_documento"))
            ws.cell(row=fila, column=2, value=gestante.get("documento"))
            ws.cell(row=fila, column=3, value=gestante.get("edad"))
            ws.cell(row=fila, column=4, value=gestante.get("municipio"))
            zona = gestante.get("zona")
            zona_desc = "Urbana" if zona == "01" else "Rural" if zona == "02" else zona
            ws.cell(row=fila, column=5, value=zona_desc)
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_distribucion_edad(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de distribución por Ruta de Vida"""
        ws = wb.create_sheet("Distribución por Ruta de Vida")
        demo = analisis.get("analisis_demografico", {})

        ws['A1'] = "DISTRIBUCIÓN DE PACIENTES POR RUTA DE VIDA"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:C1')

        ws['A2'] = "(Resolución 3280 de 2018)"
        ws['A2'].font = Font(italic=True, size=10)
        ws.merge_cells('A2:C2')

        encabezados = ["Ruta de Vida", "Cantidad", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=4, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        total_usuarios = demo.get('total_usuarios', 0)
        fila = 5
        for ruta_vida, cantidad in demo.get("distribucion_ruta_vida", {}).items():
            porcentaje = (cantidad / total_usuarios * 100) if total_usuarios > 0 else 0
            ws.cell(row=fila, column=1, value=ruta_vida)
            ws.cell(row=fila, column=2, value=cantidad)
            ws.cell(row=fila, column=3, value=f"{porcentaje:.1f}%")
            fila += 1

        # Distribución por sexo y ruta de vida
        fila += 2
        ws[f'A{fila}'] = "Distribución por Sexo y Ruta de Vida"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        encabezados = ["Sexo", "Ruta de Vida", "Cantidad"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        dist_sexo_edad = demo.get("distribucion_sexo_edad", {})
        for sexo in ["F", "M"]:
            sexo_desc = "Femenino" if sexo == "F" else "Masculino"
            for ruta_vida, cantidad in dist_sexo_edad.get(sexo, {}).items():
                ws.cell(row=fila, column=1, value=sexo_desc)
                ws.cell(row=fila, column=2, value=ruta_vida)
                ws.cell(row=fila, column=3, value=cantidad)
                fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_distribucion_territorial(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de distribución territorial"""
        ws = wb.create_sheet("Distribución Territorial")
        demo = analisis.get("analisis_demografico", {})

        ws['A1'] = "DISTRIBUCIÓN TERRITORIAL DE PACIENTES"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:C1')

        # Distribución por zona territorial
        ws['A3'] = "Distribución por Zona Territorial"
        ws['A3'].font = Font(bold=True, size=12)

        encabezados = ["Zona", "Cantidad", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=4, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        total_usuarios = demo.get('total_usuarios', 0)
        fila = 5
        for zona, cantidad in demo.get("distribucion_zona_territorial", {}).items():
            zona_desc = "Urbana" if zona == "01" else "Rural" if zona == "02" else zona
            porcentaje = (cantidad / total_usuarios * 100) if total_usuarios > 0 else 0
            ws.cell(row=fila, column=1, value=zona_desc)
            ws.cell(row=fila, column=2, value=cantidad)
            ws.cell(row=fila, column=3, value=f"{porcentaje:.1f}%")
            fila += 1

        # Distribución por municipio
        fila += 2
        ws[f'A{fila}'] = "Distribución por Municipio"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        encabezados = ["Código Municipio", "Cantidad", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for municipio, cantidad in demo.get("usuarios_por_municipio", {}).items():
            porcentaje = (cantidad / total_usuarios * 100) if total_usuarios > 0 else 0
            ws.cell(row=fila, column=1, value=municipio)
            ws.cell(row=fila, column=2, value=cantidad)
            ws.cell(row=fila, column=3, value=f"{porcentaje:.1f}%")
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_servicios(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis de servicios"""
        ws = wb.create_sheet("Análisis de Servicios")
        serv = analisis.get("analisis_servicios", {})

        ws['A1'] = "ANÁLISIS DE SERVICIOS"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:C1')

        # Resumen
        ws['A3'] = "Resumen de Servicios"
        ws['A3'].font = Font(bold=True, size=12)
        datos_resumen = [
            ["Total de Consultas", serv.get('total_consultas', 0)],
            ["Servicios sin Autorización", serv.get('servicios_sin_autorizacion', 0)],
            ["Valor Total de Servicios", f"${serv.get('total_valor_servicios', 0):,.2f}"],
            ["Total Copagos", f"${serv.get('total_copagos', 0):,.2f}"]
        ]

        fila = 4
        for campo, valor in datos_resumen:
            ws[f'A{fila}'] = campo
            ws[f'A{fila}'].font = Font(bold=True)
            ws[f'B{fila}'] = valor
            fila += 1

        # Distribución por finalidad
        fila += 2
        ws[f'A{fila}'] = "Distribución por Finalidad de Atención"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        encabezados = ["Código", "Descripción", "Cantidad"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for finalidad, cantidad in serv.get("distribucion_finalidad", {}).items():
            ws.cell(row=fila, column=1, value=finalidad)
            ws.cell(row=fila, column=2, value=self._obtener_descripcion_finalidad(finalidad))
            ws.cell(row=fila, column=3, value=cantidad)
            fila += 1

        # Distribución por causa de atención
        fila += 2
        ws[f'A{fila}'] = "Distribución por Causa de Atención"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for causa, cantidad in serv.get("distribucion_causa_atencion", {}).items():
            ws.cell(row=fila, column=1, value=causa)
            ws.cell(row=fila, column=2, value=self._obtener_descripcion_causa(causa))
            ws.cell(row=fila, column=3, value=cantidad)
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_acudientes(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis de acudientes"""
        ws = wb.create_sheet("Acudientes")
        acud = analisis.get("analisis_acudientes", {})

        ws['A1'] = "ANÁLISIS DE ACUDIENTES"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:D1')

        ws['A3'] = f"Total de Acudientes: {acud.get('total_acudientes', 0)}"
        ws['A4'] = f"Menores con Acudiente: {acud.get('total_menores_con_acudiente', 0)}"

        # Tabla de menores con acudiente
        encabezados = ["Documento Paciente", "Edad", "Sexo", "Documento Acudiente"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=6, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        fila = 7
        for menor in acud.get("menores_con_acudiente", []):
            ws.cell(row=fila, column=1, value=menor.get("paciente"))
            ws.cell(row=fila, column=2, value=menor.get("edad"))
            ws.cell(row=fila, column=3, value=menor.get("sexo", ""))
            ws.cell(row=fila, column=4, value=menor.get("acudiente"))
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_validacion_calidad(self, wb: Workbook, validacion: Dict[str, Any]):
        """Crea hoja de validación de calidad de datos"""
        ws = wb.create_sheet("Validación de Calidad")

        # Título
        ws['A1'] = "VALIDACIÓN DE CALIDAD DE DATOS - RESOLUCIÓN 2275 DE 2023"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
        ws.merge_cells('A1:F1')

        # Resumen de calidad
        ws['A3'] = "RESUMEN DE CALIDAD"
        ws['A3'].font = Font(bold=True, size=12)
        ws.merge_cells('A3:C3')

        calidad = validacion.get("calidad_datos", "N/A")
        total = validacion.get("total_anomalias", 0)

        ws['A5'] = "Nivel de Calidad:"
        ws['A5'].font = Font(bold=True)
        ws['B5'] = calidad

        # Color según calidad
        color_calidad = {
            "EXCELENTE": "00B050",
            "BUENA": "92D050",
            "REGULAR": "FFC000",
            "DEFICIENTE": "FF6600",
            "CRÍTICA": "C00000"
        }
        ws['B5'].fill = PatternFill(start_color=color_calidad.get(calidad, "FFFFFF"),
                                     end_color=color_calidad.get(calidad, "FFFFFF"),
                                     fill_type="solid")
        ws['B5'].font = Font(bold=True, color="FFFFFF")

        ws['A6'] = "Total de Anomalías:"
        ws['A6'].font = Font(bold=True)
        ws['B6'] = total

        # Distribución por severidad
        ws['A8'] = "Distribución por Severidad"
        ws['A8'].font = Font(bold=True, size=12)

        encabezados = ["Severidad", "Cantidad", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=9, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        por_severidad = validacion.get("por_severidad", {})
        fila = 10
        colores_sev = {"ALTA": "C00000", "MEDIA": "FFC000", "BAJA": "92D050"}

        for sev in ["ALTA", "MEDIA", "BAJA"]:
            cantidad = por_severidad.get(sev, 0)
            porcentaje = (cantidad / total * 100) if total > 0 else 0

            ws.cell(row=fila, column=1, value=sev)
            ws.cell(row=fila, column=1).fill = PatternFill(start_color=colores_sev[sev],
                                                            end_color=colores_sev[sev],
                                                            fill_type="solid")
            ws.cell(row=fila, column=1).font = Font(bold=True, color="FFFFFF")
            ws.cell(row=fila, column=2, value=cantidad)
            ws.cell(row=fila, column=3, value=f"{porcentaje:.1f}%")
            fila += 1

        # Top 10 tipos de anomalías
        fila += 2
        ws[f'A{fila}'] = "Top 10 Tipos de Anomalías"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        encabezados = ["Tipo de Anomalía", "Cantidad"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for item in validacion.get("top_10_anomalias", []):
            ws.cell(row=fila, column=1, value=item.get("tipo", "").replace("_", " "))
            ws.cell(row=fila, column=2, value=item.get("cantidad", 0))
            fila += 1

        # Detalle de anomalías
        fila += 2
        ws[f'A{fila}'] = "DETALLE DE ANOMALÍAS ENCONTRADAS"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{fila}:G{fila}')
        fila += 1

        encabezados = ["#", "Severidad", "Tipo", "Registro", "Campo", "Descripción", "Valor Actual", "Documento Paciente"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for idx, anomalia in enumerate(validacion.get("detalle_anomalias", []), 1):
            ws.cell(row=fila, column=1, value=idx)

            sev = anomalia.get("severidad", "MEDIA")
            cell_sev = ws.cell(row=fila, column=2, value=sev)
            cell_sev.fill = PatternFill(start_color=colores_sev.get(sev, "FFFFFF"),
                                        end_color=colores_sev.get(sev, "FFFFFF"),
                                        fill_type="solid")
            cell_sev.font = Font(bold=True, color="FFFFFF")

            ws.cell(row=fila, column=3, value=anomalia.get("tipo", "").replace("_", " "))
            ws.cell(row=fila, column=4, value=anomalia.get("registro", ""))
            ws.cell(row=fila, column=5, value=anomalia.get("campo", ""))
            ws.cell(row=fila, column=6, value=anomalia.get("descripcion", ""))
            ws.cell(row=fila, column=7, value=anomalia.get("valor_actual", ""))
            ws.cell(row=fila, column=8, value=anomalia.get("documento_paciente", ""))
            fila += 1

        # Ajustar columnas
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 35
        ws.column_dimensions['D'].width = 10
        ws.column_dimensions['E'].width = 30
        ws.column_dimensions['F'].width = 60
        ws.column_dimensions['G'].width = 20
        ws.column_dimensions['H'].width = 20

    def _crear_hoja_alertas(self, wb: Workbook, analisis_avanzado: Dict[str, Any]):
        """Crea hoja de alertas automáticas"""
        ws = wb.create_sheet("Alertas Automáticas")

        # Título
        ws['A1'] = "SISTEMA DE ALERTAS AUTOMÁTICAS"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="FF6600", end_color="FF6600", fill_type="solid")
        ws.merge_cells('A1:E1')

        alertas = analisis_avanzado.get("alertas", [])

        ws['A3'] = f"Total de Alertas Generadas: {len(alertas)}"
        ws['A3'].font = Font(bold=True, size=12)

        if alertas:
            # Distribución por nivel
            ws['A5'] = "Distribución por Nivel de Severidad"
            ws['A5'].font = Font(bold=True, size=11)

            niveles = {}
            for alerta in alertas:
                nivel = alerta.get("nivel", "MEDIA")
                niveles[nivel] = niveles.get(nivel, 0) + 1

            fila = 6
            for nivel, cantidad in niveles.items():
                ws[f'A{fila}'] = nivel
                ws[f'B{fila}'] = cantidad
                fila += 1

            # Detalle de todas las alertas
            fila += 2
            ws[f'A{fila}'] = "DETALLE DE ALERTAS"
            ws[f'A{fila}'].font = Font(bold=True, size=12)
            ws.merge_cells(f'A{fila}:E{fila}')
            fila += 1

            encabezados = ["Nivel", "Tipo", "Mensaje", "Recomendación", "Acción Requerida"]
            for col, encabezado in enumerate(encabezados, start=1):
                cell = ws.cell(row=fila, column=col, value=encabezado)
                self._aplicar_estilo_encabezado(cell)
            fila += 1

            colores_nivel = {"ALTA": "C00000", "MEDIA": "FFC000", "BAJA": "92D050"}

            for alerta in alertas:
                nivel = alerta.get("nivel", "MEDIA")

                cell_nivel = ws.cell(row=fila, column=1, value=nivel)
                cell_nivel.fill = PatternFill(start_color=colores_nivel.get(nivel, "FFFFFF"),
                                              end_color=colores_nivel.get(nivel, "FFFFFF"),
                                              fill_type="solid")
                cell_nivel.font = Font(bold=True, color="FFFFFF")

                ws.cell(row=fila, column=2, value=alerta.get("tipo", ""))
                ws.cell(row=fila, column=3, value=alerta.get("mensaje", ""))
                ws.cell(row=fila, column=4, value=alerta.get("recomendacion", ""))
                ws.cell(row=fila, column=5, value=alerta.get("accion_requerida", ""))
                fila += 1
        else:
            ws['A5'] = "No se generaron alertas. ¡Excelente desempeño!"
            ws['A5'].font = Font(bold=True, color="00B050")

        self._ajustar_columnas(ws)

    def _crear_hoja_indicadores_calidad(self, wb: Workbook, analisis_avanzado: Dict[str, Any]):
        """Crea hoja de indicadores de calidad"""
        ws = wb.create_sheet("Indicadores de Calidad")

        # Título
        ws['A1'] = "INDICADORES DE CALIDAD - RESOLUCIÓN 202 DE 2021"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="0070C0", end_color="0070C0", fill_type="solid")
        ws.merge_cells('A1:D1')

        indicadores = analisis_avanzado.get("indicadores_calidad", {})

        # Encabezados
        encabezados = ["Indicador", "Valor", "Meta", "Cumplimiento"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=3, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        fila = 4
        datos_indicadores = [
            ["Oportunidad en Atención", indicadores.get("oportunidad_atencion", 0), 95, "%"],
            ["Completitud de Diagnósticos", indicadores.get("completitud_diagnosticos", 0), 98, "%"],
            ["Cobertura de Autorizaciones", indicadores.get("cobertura_autorizaciones", 0), 100, "%"],
            ["Calidad de Registro", indicadores.get("calidad_registro", 0), 95, "%"]
        ]

        for nombre, valor, meta, unidad in datos_indicadores:
            ws.cell(row=fila, column=1, value=nombre)

            # Verificar que valor sea numérico
            if isinstance(valor, (int, float)):
                ws.cell(row=fila, column=2, value=f"{valor:.1f}{unidad}")
                cumple = "Si" if valor >= meta else "No"
                cell_cumple = ws.cell(row=fila, column=4, value=cumple)
                if valor >= meta:
                    cell_cumple.font = Font(bold=True, color="00B050")
                else:
                    cell_cumple.font = Font(bold=True, color="C00000")
            else:
                ws.cell(row=fila, column=2, value="N/A")
                ws.cell(row=fila, column=4, value="N/A")

            ws.cell(row=fila, column=3, value=f"{meta}{unidad}")
            fila += 1

        # Análisis de oportunidad
        fila += 2
        ws[f'A{fila}'] = "ANÁLISIS DE OPORTUNIDAD EN ATENCIÓN"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{fila}:D{fila}')
        fila += 1

        oportunidad = analisis_avanzado.get("analisis_oportunidad", {})
        datos_oportunidad = [
            ["Servicios Oportunos (≤24h)", oportunidad.get("servicios_oportunos", 0)],
            ["Servicios con Demora Moderada (1-3 días)", oportunidad.get("servicios_moderados", 0)],
            ["Servicios con Demora Alta (>3 días)", oportunidad.get("servicios_demorados", 0)]
        ]

        for concepto, valor in datos_oportunidad:
            ws.cell(row=fila, column=1, value=concepto)
            ws.cell(row=fila, column=2, value=valor)
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_morbilidad_especifica(self, wb: Workbook, analisis_avanzado: Dict[str, Any]):
        """Crea hoja de análisis de morbilidad específica"""
        ws = wb.create_sheet("Morbilidad Específica")

        # Título
        ws['A1'] = "ANÁLISIS DE MORBILIDAD ESPECÍFICA"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="7030A0", end_color="7030A0", fill_type="solid")
        ws.merge_cells('A1:D1')

        morbilidad = analisis_avanzado.get("morbilidad_especifica", {})

        # Enfermedades Crónicas No Transmisibles (ECNT)
        ws['A3'] = "ENFERMEDADES CRÓNICAS NO TRANSMISIBLES (ECNT)"
        ws['A3'].font = Font(bold=True, size=12)
        ws.merge_cells('A3:D3')

        encabezados = ["Patología", "Casos Identificados", "Porcentaje", "Observaciones"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=4, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)

        fila = 5
        ecnt = morbilidad.get("ecnt", {})
        total_casos = morbilidad.get("total_casos_ecnt", 0)

        for patologia, casos in ecnt.items():
            porcentaje = (casos / total_casos * 100) if total_casos > 0 else 0
            ws.cell(row=fila, column=1, value=patologia.replace("_", " ").title())
            ws.cell(row=fila, column=2, value=casos)
            ws.cell(row=fila, column=3, value=f"{porcentaje:.1f}%")

            # Añadir observación basada en cantidad
            if casos > 10:
                ws.cell(row=fila, column=4, value="Requiere seguimiento prioritario")
            elif casos > 0:
                ws.cell(row=fila, column=4, value="Requiere seguimiento")

            fila += 1

        # Eventos de Salud Pública
        fila += 2
        ws[f'A{fila}'] = "EVENTOS DE SALUD PÚBLICA IDENTIFICADOS"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{fila}:D{fila}')
        fila += 1

        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        eventos_sp = morbilidad.get("eventos_salud_publica", {})

        for evento, casos in eventos_sp.items():
            if casos > 0:
                ws.cell(row=fila, column=1, value=evento.replace("_", " ").title())
                ws.cell(row=fila, column=2, value=casos)
                ws.cell(row=fila, column=3, value="")
                ws.cell(row=fila, column=4, value="Notificar al sistema de vigilancia")
                cell_notif = ws.cell(row=fila, column=4)
                cell_notif.font = Font(bold=True, color="C00000")
                fila += 1

        if not eventos_sp or all(v == 0 for v in eventos_sp.values()):
            ws.cell(row=fila, column=1, value="No se identificaron eventos de salud pública")
            ws.cell(row=fila, column=1).font = Font(bold=True, color="00B050")

        # Grupos de riesgo
        fila += 2
        ws[f'A{fila}'] = "GRUPOS DE RIESGO IDENTIFICADOS"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{fila}:D{fila}')
        fila += 1

        grupos_riesgo = analisis_avanzado.get("grupos_riesgo", [])

        if grupos_riesgo and isinstance(grupos_riesgo, list):
            for col, encabezado in enumerate(["Grupo de Riesgo", "Casos", "Recomendación"], start=1):
                cell = ws.cell(row=fila, column=col, value=encabezado)
                self._aplicar_estilo_encabezado(cell)
            fila += 1

            for grupo in grupos_riesgo:
                if isinstance(grupo, dict):
                    ws.cell(row=fila, column=1, value=grupo.get("nombre", ""))
                    ws.cell(row=fila, column=2, value=grupo.get("cantidad", 0))
                    ws.cell(row=fila, column=3, value=grupo.get("recomendacion", ""))
                    fila += 1
                elif isinstance(grupo, str):
                    # Si es string, mostrarlo directamente
                    ws.cell(row=fila, column=1, value=grupo)
                    ws.cell(row=fila, column=2, value="N/A")
                    ws.cell(row=fila, column=3, value="")
                    fila += 1
        else:
            ws.cell(row=fila, column=1, value="No se identificaron grupos de riesgo especiales")

        self._ajustar_columnas(ws)

    def _crear_hoja_procedimientos(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis de procedimientos CUPS"""
        ws = wb.create_sheet("Procedimientos CUPS")

        proc = analisis.get("analisis_procedimientos", {})

        # Título
        ws['A1'] = "ANÁLISIS DE PROCEDIMIENTOS CUPS"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        ws.merge_cells('A1:E1')

        # Resumen
        ws['A3'] = "RESUMEN"
        ws['A3'].font = Font(bold=True, size=12)

        datos_resumen = [
            ["Total de Procedimientos", proc.get('total_procedimientos', 0)],
            ["Procedimientos Únicos (códigos CUPS)", proc.get('procedimientos_unicos', 0)],
            ["Complicaciones Reportadas", proc.get('total_complicaciones', 0)],
            ["Procedimientos con MIPRES", proc.get('procedimientos_con_mipres', 0)],
            ["Procedimientos sin Autorización", proc.get('procedimientos_sin_autorizacion', 0)],
            ["Valor Total Procedimientos", f"${proc.get('total_valor_procedimientos', 0):,.2f}"]
        ]

        fila = 4
        for campo, valor in datos_resumen:
            ws[f'A{fila}'] = campo
            ws[f'A{fila}'].font = Font(bold=True)
            ws[f'B{fila}'] = valor
            fila += 1

        # Top 20 procedimientos
        fila += 2
        ws[f'A{fila}'] = "TOP 20 PROCEDIMIENTOS MÁS FRECUENTES"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{fila}:D{fila}')
        fila += 1

        encabezados = ["Código CUPS", "Cantidad", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for proc_info in proc.get("top_20_procedimientos", []):
            ws.cell(row=fila, column=1, value=proc_info.get("codigo", ""))
            ws.cell(row=fila, column=2, value=proc_info.get("cantidad", 0))
            ws.cell(row=fila, column=3, value=f"{proc_info.get('porcentaje', 0):.1f}%")
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_tipo_usuario(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis por tipo de usuario/régimen"""
        ws = wb.create_sheet("Tipo Usuario-Régimen")

        tipo_usuario = analisis.get("analisis_tipo_usuario", {})

        # Título
        ws['A1'] = "ANÁLISIS POR TIPO DE USUARIO (RÉGIMEN)"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")
        ws.merge_cells('A1:D1')

        # Distribución
        ws['A3'] = "DISTRIBUCIÓN POR TIPO DE USUARIO"
        ws['A3'].font = Font(bold=True, size=12)
        fila = 4

        encabezados = ["Código", "Descripción", "Cantidad Usuarios", "Total Servicios", "Valor Total"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        distribucion = tipo_usuario.get("distribucion_tipo_usuario", {})
        servicios_por_tipo = tipo_usuario.get("servicios_por_tipo", {})
        valor_por_tipo = tipo_usuario.get("valor_por_tipo", {})
        descripciones = tipo_usuario.get("descripciones_tipo", {})

        for tipo, cantidad in distribucion.items():
            ws.cell(row=fila, column=1, value=tipo)
            ws.cell(row=fila, column=2, value=descripciones.get(tipo, "No especificado"))
            ws.cell(row=fila, column=3, value=cantidad)
            ws.cell(row=fila, column=4, value=servicios_por_tipo.get(tipo, 0))
            ws.cell(row=fila, column=5, value=f"${valor_por_tipo.get(tipo, 0):,.2f}")
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_incapacidades(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis de incapacidades"""
        ws = wb.create_sheet("Incapacidades")

        incap = analisis.get("analisis_incapacidades", {})

        # Título
        ws['A1'] = "ANÁLISIS DE INCAPACIDADES"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="FF6600", end_color="FF6600", fill_type="solid")
        ws.merge_cells('A1:D1')

        # Resumen
        ws['A3'] = f"Total de usuarios con incapacidad: {incap.get('total_con_incapacidad', 0)}"
        ws['A3'].font = Font(bold=True, size=12)

        ws['A4'] = f"Porcentaje: {incap.get('porcentaje_incapacidad', 0):.1f}%"
        ws['A4'].font = Font(bold=True)

        # Distribución por sexo
        fila = 6
        ws[f'A{fila}'] = "DISTRIBUCIÓN POR SEXO"
        ws[f'A{fila}'].font = Font(bold=True, size=11)
        fila += 1

        dist_sexo = incap.get("distribucion_por_sexo", {})
        for sexo, cantidad in dist_sexo.items():
            sexo_desc = "Femenino" if sexo == "F" else "Masculino" if sexo == "M" else sexo
            ws[f'A{fila}'] = sexo_desc
            ws[f'B{fila}'] = cantidad
            fila += 1

        # Distribución por edad
        fila += 1
        ws[f'A{fila}'] = "DISTRIBUCIÓN POR GRUPO DE EDAD"
        ws[f'A{fila}'].font = Font(bold=True, size=11)
        fila += 1

        dist_edad = incap.get("distribucion_por_edad", {})
        for grupo, cantidad in dist_edad.items():
            ws[f'A{fila}'] = grupo
            ws[f'B{fila}'] = cantidad
            fila += 1

        # Top diagnósticos asociados
        fila += 1
        ws[f'A{fila}'] = "TOP DIAGNÓSTICOS ASOCIADOS A INCAPACIDAD"
        ws[f'A{fila}'].font = Font(bold=True, size=11)
        ws.merge_cells(f'A{fila}:D{fila}')
        fila += 1

        encabezados = ["Código", "Nombre", "Cantidad"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for diag in incap.get("top_diagnosticos_incapacidad", []):
            ws.cell(row=fila, column=1, value=diag.get("codigo", ""))
            ws.cell(row=fila, column=2, value=diag.get("nombre", ""))
            ws.cell(row=fila, column=3, value=diag.get("cantidad", 0))
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_modalidad_diagnostico(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis de modalidad y tipo de diagnóstico"""
        ws = wb.create_sheet("Modalidad y Tipo Diag")

        modal = analisis.get("analisis_modalidad_diagnostico", {})

        # Título
        ws['A1'] = "ANÁLISIS DE MODALIDAD Y TIPO DE DIAGNÓSTICO"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="9966CC", end_color="9966CC", fill_type="solid")
        ws.merge_cells('A1:D1')

        # Tipo de diagnóstico (calidad)
        ws['A3'] = "TIPO DE DIAGNÓSTICO (CALIDAD)"
        ws['A3'].font = Font(bold=True, size=12)
        ws.merge_cells('A3:D3')
        fila = 4

        encabezados = ["Tipo", "Descripción", "Cantidad", "Porcentaje"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        tipo_diag_detalle = modal.get("tipo_diagnostico_detalle", {})
        for tipo, info in tipo_diag_detalle.items():
            ws.cell(row=fila, column=1, value=tipo)
            ws.cell(row=fila, column=2, value=info.get("descripcion", ""))
            ws.cell(row=fila, column=3, value=info.get("cantidad", 0))
            ws.cell(row=fila, column=4, value=f"{info.get('porcentaje', 0):.1f}%")
            fila += 1

        # Distribución por modalidad
        fila += 2
        ws[f'A{fila}'] = "DISTRIBUCIÓN POR MODALIDAD DE ATENCIÓN"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        fila += 1

        ws[f'A{fila}'] = "Modalidad"
        ws[f'B{fila}'] = "Cantidad"
        self._aplicar_estilo_encabezado(ws[f'A{fila}'])
        self._aplicar_estilo_encabezado(ws[f'B{fila}'])
        fila += 1

        dist_modalidad = modal.get("distribucion_modalidad", {})
        for modalidad, cantidad in dist_modalidad.items():
            ws[f'A{fila}'] = modalidad
            ws[f'B{fila}'] = cantidad
            fila += 1

        self._ajustar_columnas(ws)

    def _crear_hoja_prestadores(self, wb: Workbook, analisis: Dict[str, Any]):
        """Crea hoja de análisis de prestadores"""
        ws = wb.create_sheet("Prestadores")

        prest = analisis.get("analisis_prestadores", {})

        # Título
        ws['A1'] = "ANÁLISIS DE PRESTADORES DE SERVICIOS"
        ws['A1'].font = Font(bold=True, size=14, color="FFFFFF")
        ws['A1'].fill = PatternFill(start_color="00B0F0", end_color="00B0F0", fill_type="solid")
        ws.merge_cells('A1:E1')

        ws['A3'] = f"Total de prestadores: {prest.get('total_prestadores', 0)}"
        ws['A3'].font = Font(bold=True, size=12)

        # Top 10 prestadores
        fila = 5
        ws[f'A{fila}'] = "TOP 10 PRESTADORES"
        ws[f'A{fila}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{fila}:E{fila}')
        fila += 1

        encabezados = ["Código Prestador", "Total Servicios", "Consultas", "Procedimientos", "Valor Total"]
        for col, encabezado in enumerate(encabezados, start=1):
            cell = ws.cell(row=fila, column=col, value=encabezado)
            self._aplicar_estilo_encabezado(cell)
        fila += 1

        for prestador in prest.get("top_10_prestadores", []):
            ws.cell(row=fila, column=1, value=prestador.get("codigo_prestador", ""))
            ws.cell(row=fila, column=2, value=prestador.get("total_servicios", 0))
            ws.cell(row=fila, column=3, value=prestador.get("consultas", 0))
            ws.cell(row=fila, column=4, value=prestador.get("procedimientos", 0))
            ws.cell(row=fila, column=5, value=f"${prestador.get('valor_total', 0):,.2f}")
            fila += 1

        self._ajustar_columnas(ws)

    def _generar_grafico(self, tipo: str, titulo: str, datos: Dict[str, Any], labels: List[str] = None) -> io.BytesIO:
        """
        Genera un gráfico y lo retorna como BytesIO para insertarlo en el documento

        Args:
            tipo: Tipo de gráfico ('bar', 'pie', 'line')
            titulo: Título del gráfico
            datos: Datos para el gráfico
            labels: Etiquetas personalizadas (opcional)

        Returns:
            BytesIO con la imagen del gráfico
        """
        plt.figure(figsize=(10, 6))

        if tipo == 'bar':
            if labels:
                bars = plt.bar(labels, list(datos.values()), color='steelblue')
            else:
                bars = plt.bar(datos.keys(), datos.values(), color='steelblue')
            plt.xticks(rotation=45, ha='right')

            # Agregar etiquetas de datos en las barras
            for bar in bars:
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width()/2., height,
                        f'{int(height)}',
                        ha='center', va='bottom', fontsize=9, fontweight='bold')

        elif tipo == 'pie':
            if labels:
                wedges, texts, autotexts = plt.pie(datos.values(), labels=labels, autopct='%1.1f%%', startangle=90)
            else:
                wedges, texts, autotexts = plt.pie(datos.values(), labels=datos.keys(), autopct='%1.1f%%', startangle=90)

            # Mejorar legibilidad de las etiquetas
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(10)

            plt.axis('equal')

        elif tipo == 'line':
            if labels:
                line = plt.plot(labels, list(datos.values()), marker='o', linewidth=2, markersize=8)
            else:
                line = plt.plot(list(datos.keys()), list(datos.values()), marker='o', linewidth=2, markersize=8)
            plt.xticks(rotation=45, ha='right')

            # Agregar etiquetas de datos en los puntos
            valores = list(datos.values()) if labels else list(datos.values())
            etiquetas_x = labels if labels else list(datos.keys())
            for i, (x, y) in enumerate(zip(range(len(valores)), valores)):
                plt.text(x, y, f'{int(y)}', ha='center', va='bottom', fontsize=9, fontweight='bold')

        plt.title(titulo, fontsize=14, fontweight='bold', pad=20)
        plt.tight_layout()

        # Guardar gráfico en BytesIO
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        buffer.seek(0)
        plt.close()

        return buffer

    def generar_informe_docx(self, analisis: Dict[str, Any], datos_rips: Dict[str, Any], validacion: Dict[str, Any], nombre_archivo: str) -> str:
        """
        Genera un informe gerencial en formato DOCX con gráficos

        Args:
            analisis: Diccionario con análisis completo
            datos_rips: Datos originales RIPS
            nombre_archivo: Nombre del archivo (sin extensión)

        Returns:
            Ruta del archivo generado
        """
        ruta = self.directorio_reporte / f"{nombre_archivo}.docx"

        doc = Document()

        # Configurar estilo del documento
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Título principal
        titulo = doc.add_heading('INFORME GERENCIAL', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitulo = doc.add_heading('Análisis de Registros Individuales de Prestación de Servicios (RIPS)', level=2)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Información general
        info_gen = analisis.get("informacion_general", {})
        doc.add_paragraph(f"Fecha de Generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Número de Factura: {info_gen.get('num_factura', 'N/A')}")
        doc.add_paragraph(f"Documento Obligado: {info_gen.get('num_documento_obligado', 'N/A')}")

        doc.add_page_break()

        # 1. RESUMEN EJECUTIVO
        doc.add_heading('1. RESUMEN EJECUTIVO', level=1)

        demo = analisis.get("analisis_demografico", {})
        diag = analisis.get("analisis_diagnosticos", {})
        serv = analisis.get("analisis_servicios", {})

        resumen_text = f"""
Este informe presenta un análisis detallado de los servicios de salud prestados, basado en la información contenida en los archivos RIPS (Resolución 2275 de 2023).

Durante el periodo analizado se atendieron {demo.get('total_usuarios', 0)} usuarios, generando un total de {serv.get('total_consultas', 0)} consultas. Se identificaron {diag.get('diagnosticos_unicos', 0)} diagnósticos únicos, siendo los más frecuentes aquellos relacionados con consultas de rutina, seguimiento y atención preventiva.

El análisis demográfico muestra una distribución predominante del sexo femenino con {demo.get('distribucion_sexo', {}).get('F', 0)} pacientes ({(demo.get('distribucion_sexo', {}).get('F', 0) / demo.get('total_usuarios', 1) * 100):.1f}%), frente a {demo.get('distribucion_sexo', {}).get('M', 0)} pacientes masculinos ({(demo.get('distribucion_sexo', {}).get('M', 0) / demo.get('total_usuarios', 1) * 100):.1f}%).
"""
        doc.add_paragraph(resumen_text.strip())

        # 2. ANÁLISIS DEMOGRÁFICO
        doc.add_page_break()
        doc.add_heading('2. ANÁLISIS DEMOGRÁFICO', level=1)

        # Gráfico de distribución por sexo
        doc.add_heading('2.1 Distribución por Sexo', level=2)
        dist_sexo = demo.get("distribucion_sexo", {})
        dist_sexo_labels = {"F": "Femenino", "M": "Masculino"}
        datos_sexo = {dist_sexo_labels.get(k, k): v for k, v in dist_sexo.items()}

        grafico_sexo = self._generar_grafico('pie', 'Distribución de Pacientes por Sexo', datos_sexo)
        doc.add_picture(grafico_sexo, width=Inches(5.5))

        p = doc.add_paragraph()
        p.add_run(f"Total de Femenino: {dist_sexo.get('F', 0)}\n").bold = True
        p.add_run(f"Total de Masculino: {dist_sexo.get('M', 0)}\n").bold = True

        # Gráfico de distribución por Ruta de Vida
        doc.add_heading('2.2 Distribución por Ruta de Vida (Resolución 3280 de 2018)', level=2)
        dist_ruta_vida = demo.get("distribucion_ruta_vida", {})

        grafico_ruta_vida = self._generar_grafico('bar', 'Distribución de Pacientes por Ruta de Vida', dist_ruta_vida)
        doc.add_picture(grafico_ruta_vida, width=Inches(6))

        # Tabla de distribución por Ruta de Vida
        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = 'Light Grid Accent 1'
        encabezados = tabla.rows[0].cells
        encabezados[0].text = 'Ruta de Vida'
        encabezados[1].text = 'Cantidad'
        encabezados[2].text = 'Porcentaje'

        total_usuarios = demo.get('total_usuarios', 0)
        for ruta_vida, cantidad in dist_ruta_vida.items():
            row_cells = tabla.add_row().cells
            row_cells[0].text = ruta_vida
            row_cells[1].text = str(cantidad)
            porcentaje = (cantidad / total_usuarios * 100) if total_usuarios > 0 else 0
            row_cells[2].text = f"{porcentaje:.1f}%"

        # 3. ANÁLISIS TERRITORIAL
        doc.add_page_break()
        doc.add_heading('3. DISTRIBUCIÓN TERRITORIAL', level=1)

        # Gráfico de distribución por zona territorial
        dist_zona = demo.get("distribucion_zona_territorial", {})
        datos_zona = {}
        for zona, cantidad in dist_zona.items():
            zona_desc = "Urbana" if zona == "01" else "Rural" if zona == "02" else zona
            datos_zona[zona_desc] = cantidad

        grafico_zona = self._generar_grafico('pie', 'Distribución por Zona Territorial', datos_zona)
        doc.add_picture(grafico_zona, width=Inches(5.5))

        p = doc.add_paragraph()
        for zona_desc, cantidad in datos_zona.items():
            porcentaje = (cantidad / total_usuarios * 100) if total_usuarios > 0 else 0
            p.add_run(f"{zona_desc}: {cantidad} ({porcentaje:.1f}%)\n").bold = True

        # 4. ANÁLISIS DE DIAGNÓSTICOS
        doc.add_page_break()
        doc.add_heading('4. ANÁLISIS DE DIAGNÓSTICOS', level=1)

        doc.add_paragraph(f"Total de diagnósticos registrados: {diag.get('total_diagnosticos', 0)}")
        doc.add_paragraph(f"Diagnósticos únicos identificados: {diag.get('diagnosticos_unicos', 0)}")
        doc.add_paragraph(f"Casos con comorbilidades: {diag.get('diagnosticos_con_comorbilidades', 0)}")

        # Todos los diagnósticos
        doc.add_heading('4.1 Todos los Diagnósticos Identificados', level=2)

        todos_diag = diag.get("todos_diagnosticos", [])

        # Mostrar gráfico con los top 15 más frecuentes para mejor visualización
        top_15_diag = todos_diag[:15]
        datos_top_diag = {d['codigo']: d['cantidad'] for d in top_15_diag}

        grafico_diag = self._generar_grafico('bar', 'Top 15 Diagnósticos Más Frecuentes', datos_top_diag)
        doc.add_picture(grafico_diag, width=Inches(6))

        # Tabla de TODOS los diagnósticos
        tabla_diag = doc.add_table(rows=1, cols=4)
        tabla_diag.style = 'Light Grid Accent 1'
        encabezados_diag = tabla_diag.rows[0].cells
        encabezados_diag[0].text = 'Código'
        encabezados_diag[1].text = 'Diagnóstico'
        encabezados_diag[2].text = 'Cantidad'
        encabezados_diag[3].text = 'Porcentaje'

        for d in todos_diag:
            row_cells = tabla_diag.add_row().cells
            row_cells[0].text = d.get('codigo', '')
            row_cells[1].text = d.get('nombre', '')[:50]  # Limitar longitud para mejor visualización
            row_cells[2].text = str(d.get('cantidad', 0))
            row_cells[3].text = f"{d.get('porcentaje', 0):.1f}%"

        # 4.2 Análisis Detallado de Comorbilidades
        doc.add_heading('4.2 Análisis Detallado de Comorbilidades', level=2)

        comorbilidad_detalle = diag.get("analisis_comorbilidad_detallado", {})

        doc.add_paragraph(f"Total de pacientes con comorbilidades: {comorbilidad_detalle.get('total_pacientes_con_comorbilidades', 0)}")
        doc.add_paragraph(f"Total de comorbilidades registradas: {comorbilidad_detalle.get('total_comorbilidades_registradas', 0)}")
        doc.add_paragraph(f"Promedio de comorbilidades por paciente: {comorbilidad_detalle.get('promedio_comorbilidades_por_paciente', 0)}")

        # Diagnóstico principal con más comorbilidades
        diag_mas_comorbilidades = comorbilidad_detalle.get("diagnostico_principal_con_mas_comorbilidades")
        if diag_mas_comorbilidades:
            p = doc.add_paragraph()
            p.add_run("Diagnóstico principal con más comorbilidades asociadas: ").bold = True
            p.add_run(f"[{diag_mas_comorbilidades.get('codigo')}] {diag_mas_comorbilidades.get('nombre')} - ")
            p.add_run(f"{diag_mas_comorbilidades.get('cantidad_pacientes')} pacientes")

        # Combinaciones más frecuentes
        combinaciones_frecuentes = comorbilidad_detalle.get("combinaciones_mas_frecuentes", [])
        if combinaciones_frecuentes:
            doc.add_paragraph().add_run('Combinaciones de diagnósticos más frecuentes:').bold = True
            for i, comb in enumerate(combinaciones_frecuentes[:5], 1):
                doc.add_paragraph(f"  {i}. {comb.get('combinacion')} - {comb.get('frecuencia')} casos", style='List Number')

        # 5. ANÁLISIS DE POBLACIÓN GESTANTE
        doc.add_page_break()
        doc.add_heading('5. ANÁLISIS DE POBLACIÓN GESTANTE', level=1)

        gestantes_data = analisis.get("analisis_poblacion_gestante", {})
        total_gestantes = gestantes_data.get('total_gestantes', 0)

        if total_gestantes > 0:
            # Resumen
            doc.add_paragraph(f"Total de gestantes identificadas: {total_gestantes}")
            doc.add_paragraph(f"Edad promedio: {gestantes_data.get('edad_promedio', 0)} años")
            doc.add_paragraph(f"Gestantes adolescentes (10-17 años): {gestantes_data.get('gestantes_adolescentes', 0)}")
            doc.add_paragraph(f"Gestantes de alto riesgo: {gestantes_data.get('gestantes_alto_riesgo', 0)} ({gestantes_data.get('porcentaje_alto_riesgo', 0):.1f}%)")
            doc.add_paragraph(f"Gestantes con complicaciones: {gestantes_data.get('gestantes_con_complicaciones', 0)} ({gestantes_data.get('porcentaje_complicaciones', 0):.1f}%)")
            doc.add_paragraph(f"Cobertura de controles prenatales: {gestantes_data.get('cobertura_controles_prenatales', 0):.1f}%")

            # Controles Prenatales
            doc.add_heading('5.1 Controles Prenatales', level=2)

            dist_controles = gestantes_data.get("distribucion_controles_prenatales", {})
            if dist_controles:
                tabla_controles = doc.add_table(rows=1, cols=3)
                tabla_controles.style = 'Light Grid Accent 1'
                enc_controles = tabla_controles.rows[0].cells
                enc_controles[0].text = 'Número de Controles'
                enc_controles[1].text = 'Cantidad'
                enc_controles[2].text = 'Porcentaje'

                for num_controles, cantidad in sorted(dist_controles.items()):
                    porcentaje = (cantidad / total_gestantes * 100) if total_gestantes > 0 else 0
                    row_cells = tabla_controles.add_row().cells
                    row_cells[0].text = f"{num_controles} controles"
                    row_cells[1].text = str(cantidad)
                    row_cells[2].text = f"{porcentaje:.1f}%"

            # Distribución por edad
            doc.add_heading('5.2 Distribución por Grupo de Edad', level=2)

            dist_edad = gestantes_data.get("distribucion_edad_gestantes", {})
            datos_edad_gest = {k: v for k, v in dist_edad.items() if v > 0}

            if datos_edad_gest:
                grafico_gest = self._generar_grafico('pie', 'Distribución de Gestantes por Grupo de Edad', datos_edad_gest)
                doc.add_picture(grafico_gest, width=Inches(5.5))

            # Tabla de distribución
            tabla_gest = doc.add_table(rows=1, cols=3)
            tabla_gest.style = 'Light Grid Accent 1'
            enc_gest = tabla_gest.rows[0].cells
            enc_gest[0].text = 'Grupo de Edad'
            enc_gest[1].text = 'Cantidad'
            enc_gest[2].text = 'Porcentaje'

            for grupo, cantidad in dist_edad.items():
                if cantidad > 0:
                    porcentaje = (cantidad / total_gestantes * 100) if total_gestantes > 0 else 0
                    row_cells = tabla_gest.add_row().cells
                    row_cells[0].text = grupo
                    row_cells[1].text = str(cantidad)
                    row_cells[2].text = f"{porcentaje:.1f}%"

            # Top diagnósticos en gestantes
            doc.add_heading('5.3 Diagnósticos Principales en Población Gestante', level=2)

            top_diag_gest = gestantes_data.get("top_diagnosticos_gestantes", [])
            if top_diag_gest:
                for i, diag_info in enumerate(top_diag_gest[:5], 1):
                    doc.add_paragraph(f"  {i}. [{diag_info.get('codigo')}] {diag_info.get('nombre')} - {diag_info.get('cantidad')} casos", style='List Number')
        else:
            doc.add_paragraph("No se identificaron pacientes gestantes en el período analizado.")

        # 6. ANÁLISIS DE MOTIVOS DE CONSULTA
        doc.add_page_break()
        doc.add_heading('6. MOTIVOS DE CONSULTA', level=1)

        # Distribución por causa de atención
        dist_causa = serv.get("distribucion_causa_atencion", {})
        datos_causa = {}
        for causa, cantidad in dist_causa.items():
            desc = self._obtener_descripcion_causa(causa)
            datos_causa[f"{causa}"] = cantidad

        grafico_causa = self._generar_grafico('bar', 'Distribución por Causa de Atención', datos_causa)
        doc.add_picture(grafico_causa, width=Inches(6))

        # Tabla de causas
        tabla_causa = doc.add_table(rows=1, cols=3)
        tabla_causa.style = 'Light Grid Accent 1'
        enc_causa = tabla_causa.rows[0].cells
        enc_causa[0].text = 'Código'
        enc_causa[1].text = 'Descripción'
        enc_causa[2].text = 'Cantidad'

        for causa, cantidad in dist_causa.items():
            row_cells = tabla_causa.add_row().cells
            row_cells[0].text = causa
            row_cells[1].text = self._obtener_descripcion_causa(causa)
            row_cells[2].text = str(cantidad)

        # 7. ANÁLISIS DE SERVICIOS
        doc.add_page_break()
        doc.add_heading('7. ANÁLISIS DE SERVICIOS', level=1)

        doc.add_paragraph(f"Total de consultas realizadas: {serv.get('total_consultas', 0)}")
        doc.add_paragraph(f"Servicios sin autorización: {serv.get('servicios_sin_autorizacion', 0)}")
        doc.add_paragraph(f"Valor total de servicios: ${serv.get('total_valor_servicios', 0):,.2f}")
        doc.add_paragraph(f"Total de copagos: ${serv.get('total_copagos', 0):,.2f}")

        # Distribución por finalidad
        dist_finalidad = serv.get("distribucion_finalidad", {})
        datos_finalidad = {}
        for finalidad, cantidad in dist_finalidad.items():
            desc = self._obtener_descripcion_finalidad(finalidad)[:30]  # Limitar longitud
            datos_finalidad[finalidad] = cantidad

        grafico_finalidad = self._generar_grafico('pie', 'Distribución por Finalidad de Atención', datos_finalidad)
        doc.add_picture(grafico_finalidad, width=Inches(5.5))

        # 8. VALIDACIÓN DE CALIDAD DE DATOS
        doc.add_page_break()
        doc.add_heading('8. VALIDACIÓN DE CALIDAD DE DATOS', level=1)

        calidad = validacion.get("calidad_datos", "N/A")
        total_anomalias = validacion.get("total_anomalias", 0)

        p_calidad = doc.add_paragraph()
        p_calidad.add_run(f"Nivel de Calidad: ").bold = True
        run_calidad = p_calidad.add_run(calidad)
        run_calidad.bold = True
        run_calidad.font.size = Pt(14)

        # Color según calidad
        if calidad == "EXCELENTE":
            run_calidad.font.color.rgb = RGBColor(0, 176, 80)
        elif calidad == "BUENA":
            run_calidad.font.color.rgb = RGBColor(146, 208, 80)
        elif calidad == "REGULAR":
            run_calidad.font.color.rgb = RGBColor(255, 192, 0)
        elif calidad == "DEFICIENTE":
            run_calidad.font.color.rgb = RGBColor(255, 102, 0)
        else:  # CRÍTICA
            run_calidad.font.color.rgb = RGBColor(192, 0, 0)

        doc.add_paragraph(f"Total de anomalías detectadas: {total_anomalias}")

        if total_anomalias > 0:
            # Distribución por severidad
            doc.add_heading('8.1 Distribución por Severidad', level=2)

            por_severidad = validacion.get("por_severidad", {})
            tabla_sev = doc.add_table(rows=1, cols=3)
            tabla_sev.style = 'Light Grid Accent 1'
            enc_sev = tabla_sev.rows[0].cells
            enc_sev[0].text = 'Severidad'
            enc_sev[1].text = 'Cantidad'
            enc_sev[2].text = 'Porcentaje'

            for sev in ["ALTA", "MEDIA", "BAJA"]:
                cantidad = por_severidad.get(sev, 0)
                porcentaje = (cantidad / total_anomalias * 100) if total_anomalias > 0 else 0
                row_cells = tabla_sev.add_row().cells
                row_cells[0].text = sev
                row_cells[1].text = str(cantidad)
                row_cells[2].text = f"{porcentaje:.1f}%"

            # Top 10 anomalías
            doc.add_heading('8.2 Tipos de Anomalías Más Frecuentes', level=2)

            top_anomalias = validacion.get("top_10_anomalias", [])[:10]
            if top_anomalias:
                for item in top_anomalias:
                    tipo = item.get("tipo", "").replace("_", " ")
                    cantidad = item.get("cantidad", 0)
                    doc.add_paragraph(f"  • {tipo}: {cantidad} casos", style='List Bullet')

            # Detalle de anomalías críticas
            doc.add_heading('8.3 Detalle de Anomalías Críticas', level=2)

            anomalias_criticas = [a for a in validacion.get("detalle_anomalias", []) if a.get("severidad") == "ALTA"][:20]
            if anomalias_criticas:
                tabla_anom = doc.add_table(rows=1, cols=5)
                tabla_anom.style = 'Light Grid Accent 1'
                enc_anom = tabla_anom.rows[0].cells
                enc_anom[0].text = 'Reg.'
                enc_anom[1].text = 'Tipo'
                enc_anom[2].text = 'Campo'
                enc_anom[3].text = 'Descripción'
                enc_anom[4].text = 'Documento'

                for anomalia in anomalias_criticas:
                    row_cells = tabla_anom.add_row().cells
                    row_cells[0].text = str(anomalia.get("registro", ""))
                    row_cells[1].text = anomalia.get("tipo", "").replace("_", " ")[:30]
                    row_cells[2].text = anomalia.get("campo", "")[:20]
                    row_cells[3].text = anomalia.get("descripcion", "")[:60]
                    row_cells[4].text = str(anomalia.get("documento_paciente", ""))

                doc.add_paragraph(f"\nNota: Se muestran las primeras 20 anomalías de severidad ALTA. Total de anomalías ALTA: {por_severidad.get('ALTA', 0)}")
            else:
                doc.add_paragraph("No se encontraron anomalías de severidad ALTA.")
        else:
            doc.add_paragraph("¡Excelente! No se encontraron anomalías en los datos. La información cumple con todos los estándares de calidad de la Resolución 2275 de 2023.")

        # 9. CONCLUSIONES Y RECOMENDACIONES
        doc.add_page_break()
        doc.add_heading('9. CONCLUSIONES Y RECOMENDACIONES', level=1)

        conclusiones = f"""
Con base en el análisis realizado, se pueden destacar las siguientes conclusiones:

1. Perfil Demográfico: La población atendida presenta una distribución predominante del sexo femenino ({(demo.get('distribucion_sexo', {}).get('F', 0) / total_usuarios * 100):.1f}%), lo cual puede estar asociado a servicios de salud reproductiva, control prenatal y atención materno-infantil.

2. Distribución Territorial: La mayor parte de la población atendida ({(list(datos_zona.values())[0] / total_usuarios * 100):.1f}%) proviene de zona {list(datos_zona.keys())[0].lower()}, lo que permite focalizar estrategias de atención según las características de esta población.

3. Diagnósticos Principales: Los diagnósticos más frecuentes corresponden a consultas de rutina, seguimiento y atención preventiva, reflejando una orientación hacia la promoción y prevención en salud.

4. Servicios Prestados: Se observa que {serv.get('servicios_sin_autorizacion', 0)} servicios no cuentan con autorización previa, lo que representa un área de mejora en los procesos administrativos.

Recomendaciones:

• Fortalecer los programas de promoción y prevención, especialmente aquellos orientados a la población femenina.
• Mejorar los procesos de autorización de servicios para reducir el número de servicios sin autorización previa.
• Implementar estrategias diferenciadas según la zona territorial de origen de los pacientes.
• Continuar con el seguimiento y análisis periódico de los indicadores de salud para la toma de decisiones basada en evidencia.
"""
        doc.add_paragraph(conclusiones.strip())

        # Pie de página
        doc.add_page_break()
        p_final = doc.add_paragraph()
        p_final.add_run('Fin del Informe Gerencial\n').bold = True
        p_final.add_run(f'Generado automáticamente el {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(ruta)
        return str(ruta)
